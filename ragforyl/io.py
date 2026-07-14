from __future__ import annotations

import hashlib
import json
import re
from collections.abc import Iterable
from pathlib import Path
from typing import Any

from docx import Document
from pypdf import PdfReader

from ragforyl.config import SUPPORTED_EXTENSIONS
from ragforyl.models import SourceDocument


class SourceError(RuntimeError):
    pass


def sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def sha256_text(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def safe_filename(name: str) -> str:
    raw = Path(name or "document.txt").name
    stem = re.sub(r"[^\w\-.\u4e00-\u9fff]+", "_", Path(raw).stem).strip("._") or "document"
    suffix = Path(raw).suffix.lower()
    return f"{stem[:100]}{suffix}"


def read_text_file(path: Path) -> str:
    data = path.read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise SourceError(f"Cannot decode text file: {path.name}")


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md", ".markdown"}:
        return read_text_file(path)
    if suffix == ".pdf":
        reader = PdfReader(str(path))
        pages = []
        empty_pages = []
        for index, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if text:
                pages.append(f"[Page {index}]\n{text}")
            else:
                empty_pages.append(index)
        if empty_pages:
            try:
                ocr_pages = _ocr_pdf_pages(path, empty_pages)
            except SourceError:
                if not pages:
                    raise
            else:
                pages.extend(f"[Page {index}]\n{text}" for index, text in ocr_pages if text)
                pages.sort(key=_page_number)
        return "\n\n".join(pages)
    if suffix == ".docx":
        document = Document(str(path))
        paragraphs = [
            paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()
        ]
        for table in document.tables:
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells]
                if any(values):
                    paragraphs.append(" | ".join(values))
        return "\n\n".join(paragraphs)
    raise SourceError(f"Unsupported file type: {path.suffix}")


def _ocr_pdf_pages(path: Path, page_numbers: list[int]) -> list[tuple[int, str]]:
    try:
        import fitz
        from rapidocr_onnxruntime import RapidOCR
    except ImportError as exc:
        raise SourceError(
            f"{path.name} contains scanned pages. Install OCR support with: pip install -e .[ocr]"
        ) from exc
    engine = RapidOCR()
    document = fitz.open(str(path))
    results = []
    try:
        for page_number in page_numbers:
            page = document.load_page(page_number - 1)
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            detected, _ = engine(pixmap.tobytes("png"))
            text = "\n".join(
                str(item[1]).strip() for item in detected or [] if str(item[1]).strip()
            )
            results.append((page_number, text))
    finally:
        document.close()
    return results


def _page_number(page_text: str) -> int:
    match = re.match(r"\[Page (\d+)]", page_text)
    return int(match.group(1)) if match else 0


def load_source_documents(source_dir: Path) -> list[SourceDocument]:
    if not source_dir.exists():
        raise SourceError(f"Source directory does not exist: {source_dir}")
    paths = sorted(
        path
        for path in source_dir.rglob("*")
        if path.is_file()
        and path.suffix.lower() in SUPPORTED_EXTENSIONS
        and not path.name.startswith(".")
    )
    if not paths:
        raise SourceError(
            "No supported documents found. "
            "Add .txt, .md, .pdf, or .docx files to the source directory."
        )

    documents = []
    seen_hashes = set()
    for path in paths:
        binary = path.read_bytes()
        digest = sha256_bytes(binary)
        if digest in seen_hashes:
            continue
        text = extract_text(path).replace("\x00", "").strip()
        if not text:
            continue
        relative_path = path.relative_to(source_dir).as_posix()
        source_id = f"SRC_{hashlib.sha1((relative_path + digest).encode('utf-8')).hexdigest()[:12]}"
        documents.append(
            SourceDocument(
                id=source_id,
                title=path.stem,
                relative_path=relative_path,
                sha256=digest,
                text=text,
                size_bytes=len(binary),
            )
        )
        seen_hashes.add(digest)
    if not documents:
        raise SourceError("Documents were found but no readable text could be extracted.")
    return documents


def write_json(path: Path, payload: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_suffix(path.suffix + ".tmp")
    temporary.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    temporary.replace(path)


def write_jsonl(path: Path, rows: Iterable[dict[str, Any]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_suffix(path.suffix + ".tmp")
    with temporary.open("w", encoding="utf-8", newline="\n") as handle:
        for row in rows:
            handle.write(json.dumps(row, ensure_ascii=False) + "\n")
    temporary.replace(path)


def read_json(path: Path) -> Any:
    return json.loads(path.read_text(encoding="utf-8"))


def read_jsonl(path: Path) -> list[dict[str, Any]]:
    rows = []
    with path.open("r", encoding="utf-8") as handle:
        for line in handle:
            if line.strip():
                rows.append(json.loads(line))
    return rows
